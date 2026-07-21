import os
from datetime import date, datetime
from decimal import Decimal

from docx import Document

from app.routers.reports import generate_act_docx


class FakeUser:
    full_name = "Иванов И.И."
    username = "ivanov"
    role = "comendant"


class FakeBuilding:
    number = "5"
    name = "Главный"


class FakeService:
    name = "Ремонт кровли"
    unit = "м2"


class FakeWorkService:
    service = FakeService()
    quantity = Decimal("10.5")
    unit_price = Decimal("1200.00")
    total_price = Decimal("12600.00")


class FakeRequest:
    id = 42
    creator = FakeUser()
    created_at = datetime(2026, 7, 21, 10, 30)


class FakeMaterial:
    name = "Тройник"
    unit = "шт"


class FakeWorkMaterial:
    material = FakeMaterial()
    quantity = Decimal("2")
    unit_price = Decimal("50.00")
    total_price = Decimal("100.00")


class FakeWork:
    work_date = date(2026, 6, 15)
    building = FakeBuilding()
    work_services = [FakeWorkService()]
    total_price = Decimal("12600.00")
    request_id = 42
    request = FakeRequest()
    work_materials = [FakeWorkMaterial()]


def test_act_table_has_work_date_column():
    output = generate_act_docx([FakeWork()])
    doc = Document(output)
    tables = [t for t in doc.tables if any(cell.text == "Наименование работ" for cell in t.rows[0].cells)]
    assert tables, "Таблица работ не найдена"
    header = [cell.text for cell in tables[0].rows[0].cells]
    assert "Дата работ" in header, f"Столбец 'Дата работ' отсутствует: {header}"


def test_act_table_has_request_id_column():
    output = generate_act_docx([FakeWork()])
    doc = Document(output)
    tables = [t for t in doc.tables if any(cell.text == "Наименование работ" for cell in t.rows[0].cells)]
    assert tables, "Таблица работ не найдена"
    header = [cell.text for cell in tables[0].rows[0].cells]
    assert "№ заявки" in header, f"Столбец '№ заявки' отсутствует: {header}"
    row_values = [cell.text for cell in tables[0].rows[1].cells]
    assert row_values[1] == "42 от 21.07.26", f"Номер заявки с датой не выведен: {row_values}"


def test_act_includes_comendant_name():
    output = generate_act_docx([FakeWork()])
    doc = Document(output)
    found = any("Комендант:" in p.text and "Иванов И.И." in p.text for p in doc.paragraphs)
    assert found, "ФИО коменданта не найдено в акте"


def test_act_table_request_id_is_blank_for_unlinked_work():
    class UnlinkedFakeWork(FakeWork):
        request_id = None
        request = None

    output = generate_act_docx([UnlinkedFakeWork()])
    doc = Document(output)
    tables = [t for t in doc.tables if any(cell.text == "Наименование работ" for cell in t.rows[0].cells)]
    row_values = [cell.text for cell in tables[0].rows[1].cells]
    assert row_values[1] == "", f"Поле № заявки должно быть пустым: {row_values}"


def test_act_header_date_is_russian():
    russian_months = [
        "января", "февраля", "марта", "апреля", "мая", "июня",
        "июля", "августа", "сентября", "октября", "ноября", "декабря"
    ]
    output = generate_act_docx([FakeWork()])
    doc = Document(output)
    found = False
    for p in doc.paragraphs:
        if "г. Пятигорск" in p.text and any(month in p.text for month in russian_months):
            found = True
            break
    assert found, "Русская дата в шапке акта не найдена"


def test_act_title_has_no_word_form():
    output = generate_act_docx([FakeWork()])
    doc = Document(output)
    first_paragraph = doc.paragraphs[0].text
    assert first_paragraph == "Акт сдачи-приемки оказанных услуг"
    assert "Форма" not in first_paragraph


def test_act_materials_table_has_request_and_date_columns():
    output = generate_act_docx([FakeWork()])
    doc = Document(output)
    tables = [t for t in doc.tables if any(cell.text == "Наименование материалов" for cell in t.rows[0].cells)]
    assert tables, "Таблица материалов не найдена"
    header = [cell.text for cell in tables[0].rows[0].cells]
    assert "№ заявки" in header, f"Столбец '№ заявки' отсутствует в материалах: {header}"
    assert "Дата работ" in header, f"Столбец 'Дата работ' отсутствует в материалах: {header}"
    row_values = [cell.text for cell in tables[0].rows[1].cells]
    assert row_values[1] == "42 от 21.07.26", f"Номер заявки с датой в материалах не выведен: {row_values}"
    assert row_values[2] == "15.06.2026", f"Дата работ в материалах не выведена: {row_values}"
