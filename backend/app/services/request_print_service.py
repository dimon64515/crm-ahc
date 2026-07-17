import io
import os
import re
from datetime import datetime
from typing import Optional

from docx import Document

from app.core.config import get_settings
from app.models import Request


REQUEST_TEMPLATE_FILENAME = "Заявка принятая.docx"


def _get_template_path() -> str:
    settings = get_settings()
    return os.path.join(settings.UPLOAD_DIR, REQUEST_TEMPLATE_FILENAME)


def _format_date(d: Optional[datetime]) -> str:
    if not d:
        return ""
    return d.strftime("%d.%m.%Y")


def _replace_in_paragraph(paragraph, old_text: str, new_text: str) -> bool:
    """Заменяет подстроку в абзаце, даже если текст разбит на несколько run."""
    full_text = paragraph.text
    if old_text not in full_text:
        return False

    new_full_text = full_text.replace(old_text, new_text, 1)
    first_run = paragraph.runs[0] if paragraph.runs else None
    font = first_run.font if first_run else None

    paragraph.clear()
    run = paragraph.add_run(new_full_text)
    if font:
        run.font.name = font.name
        run.font.size = font.size
        run.bold = font.bold
        run.font.italic = font.italic
    return True


def _replace_first_in_doc(doc, old_text: str, new_text: str) -> bool:
    """Заменяет первое вхождение подстроки в документе (абзацы, затем таблицы)."""
    for paragraph in doc.paragraphs:
        if old_text in paragraph.text:
            _replace_in_paragraph(paragraph, old_text, new_text)
            return True
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if old_text in paragraph.text:
                        _replace_in_paragraph(paragraph, old_text, new_text)
                        return True
    return False


def _replace_all_in_doc(doc, old_text: str, new_text: str) -> int:
    """Заменяет все вхождения подстроки в документе (абзацы и таблицы)."""
    count = 0
    for paragraph in doc.paragraphs:
        if old_text in paragraph.text:
            _replace_in_paragraph(paragraph, old_text, new_text)
            count += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if old_text in paragraph.text:
                        _replace_in_paragraph(paragraph, old_text, new_text)
                        count += 1
    return count


def _set_cell_text(cell, text: str):
    """Очищает ячейку таблицы и записывает новый текст."""
    cell.text = text


def fill_request_template(request: Request) -> io.BytesIO:
    """Заполняет шаблон заявки данными из модели Request и возвращает DOCX в BytesIO."""
    template_path = _get_template_path()
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Шаблон заявки не найден: {template_path}")

    doc = Document(template_path)

    request_id = str(request.id) if request.id else ""
    building_number = request.building.number if request.building else ""
    creator_name = request.creator.full_name or request.creator.username if request.creator else ""
    executor_name = request.executor.full_name or request.executor.username if request.executor else ""
    description = request.description or ""
    created_date = _format_date(request.created_at)

    # Номер заявки: "Поступила ЗАЯВКА    №"
    _replace_first_in_doc(doc, "Поступила ЗАЯВКА    №", f"Поступила ЗАЯВКА    № {request_id}")

    # Заявитель: "от Комендант ..." — заменяем всю строку
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("от Комендант"):
            _replace_in_paragraph(paragraph, text, f"от Комендант {creator_name}" if creator_name else "от Комендант")
            break

    # Причина отказа оставляем пустой
    _replace_first_in_doc(doc, "Причина отказа:", "Причина отказа: —")

    # Необходимый ремонт — описание заявки
    _replace_first_in_doc(
        doc,
        "Необходим ремонт: _____________________________________________________",
        f"Необходим ремонт: {description}",
    )

    # Подпись заявителя (коменданта) в блоке "Поданная мною заявка исполнена"
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("__________ Комендант"):
            _replace_in_paragraph(
                paragraph,
                text,
                f"__________ Комендант {creator_name if creator_name else '________________________'} ________________________"
            )
            break

    # Таблица заявки
    if doc.tables:
        table = doc.tables[0]
        if len(table.rows) > 1:
            row = table.rows[1]
            # Шаблон содержит 4 колонки: Дата | Номер корпуса | ВЫЯВЛЕНО | ВЫЯВЛЕНО
            if len(row.cells) >= 4:
                _set_cell_text(row.cells[0], created_date)
                _set_cell_text(row.cells[1], building_number)
                _set_cell_text(row.cells[2], description)
                _set_cell_text(row.cells[3], description)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output
