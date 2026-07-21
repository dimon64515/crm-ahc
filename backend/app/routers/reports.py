import os
from io import BytesIO
from datetime import datetime, date
from decimal import Decimal
from typing import List

import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from openpyxl.drawing.image import Image as XLImage
from openpyxl.utils import get_column_letter
from sqlalchemy.orm import Session, joinedload

from app.database import get_db
from app.models import Work, Building, Service, User, WorkMaterial, WorkService, Request
from app.schemas import SummaryReportItem
from app.core.dependencies import get_current_user, require_director
from app.core.config import get_settings

router = APIRouter(prefix="/reports", tags=["reports"])


@router.get("/summary")
def get_summary_report(
    date_from: str = None,
    date_to: str = None,
    group_by: str = "building",
    building_id: str = None,
    db: Session = Depends(get_db),
    current_user = Depends(require_director)
):
    query = db.query(Work)
    
    if date_from:
        query = query.filter(Work.work_date >= date_from)
    if date_to:
        query = query.filter(Work.work_date <= date_to)
    if building_id:
        query = query.filter(Work.building_id == int(building_id))
    
    works = query.all()

    # Группировка
    groups = {}
    for work in works:
        if group_by == "building":
            key = work.building.number
            name = f"Корпус {work.building.number}"
        elif group_by == "contractor":
            key = work.user.full_name or work.user.username
            name = key
        else:
            key = work.work_date.strftime("%Y-%m-%d")
            name = key

        if group_by != "service":
            if key not in groups:
                groups[key] = {
                    "group_key": key,
                    "group_name": name,
                    "works_count": 0,
                    "service_total": Decimal('0'),
                    "materials_total": Decimal('0'),
                    "total": Decimal('0'),
                }
            groups[key]["works_count"] += 1
            groups[key]["service_total"] += sum((ws.total_price or Decimal('0')) for ws in work.work_services)
            groups[key]["materials_total"] += work.materials_total_price or 0
            groups[key]["total"] += work.total_price or 0
        else:
            # Группировка по услугам: каждая услуга работы — отдельная строка
            for ws in work.work_services:
                svc_key = ws.service.name
                svc_name = ws.service.name
                if svc_key not in groups:
                    groups[svc_key] = {
                        "group_key": svc_key,
                        "group_name": svc_name,
                        "works_count": 0,
                        "service_total": Decimal('0'),
                        "materials_total": Decimal('0'),
                        "total": Decimal('0'),
                    }
                groups[svc_key]["works_count"] += 1
                groups[svc_key]["service_total"] += ws.total_price or 0
    
    items = list(groups.values())
    totals = {
        "works_count": sum(i["works_count"] for i in items),
        "service_total": sum(i["service_total"] for i in items),
        "materials_total": sum(i["materials_total"] for i in items),
        "total": sum(i["total"] for i in items),
    }
    
    return {
        "group_by": group_by,
        "items": items,
        "totals": totals,
    }


@router.get("/export")
def export_works(
    date_from: str = None,
    date_to: str = None,
    building_id: str = None,
    service_id: str = None,
    user_id: str = None,
    db: Session = Depends(get_db),
    current_user = Depends(require_director)
):
    query = db.query(Work)
    
    if date_from:
        query = query.filter(Work.work_date >= date_from)
    if date_to:
        query = query.filter(Work.work_date <= date_to)
    if building_id:
        query = query.filter(Work.building_id == int(building_id))
    if service_id:
        query = query.join(WorkService).filter(WorkService.service_id == int(service_id))
    if user_id:
        query = query.filter(Work.user_id == int(user_id))

    works = query.order_by(Work.work_date.desc()).all()

    settings = get_settings()

    # Определяем максимальное количество фото и материалов
    max_photos = 0
    max_materials = 0
    for work in works:
        max_photos = max(max_photos, len(work.photos or []))
        max_materials = max(max_materials, len(work.work_materials or []))

    # Формирование DataFrame — каждая услуга работы = 1 строка
    base_columns = [
        "Дата", "Корпус", "Подрядчик", "Вид работы", "Наименование",
        "Ед.изм.", "Кол-во работы", "Цена работы", "Сумма работы",
        "Сумма материалов", "ИТОГО"
    ]
    material_columns = []
    for i in range(max_materials):
        material_columns.extend([
            f"Материал {i+1}", f"Кол-во {i+1}", f"Цена {i+1}", f"Сумма {i+1}"
        ])
    photo_columns = [f"Фото {i+1}" for i in range(max_photos)] if max_photos else []
    all_columns = base_columns + material_columns + photo_columns

    data = []
    photo_map = []  # [(excel_row_idx, col_idx, photo_file_path), ...]

    for work in works:
        work_services = list(work.work_services or [])
        if not work_services:
            # Защита от некорректных данных
            work_services = [None]

        for svc_idx, ws_item in enumerate(work_services):
            row = {
                "Дата": work.work_date,
                "Корпус": work.building.number,
                "Подрядчик": work.user.full_name or work.user.username,
                "Вид работы": ws_item.service.name if ws_item else "—",
                "Наименование": work.description if svc_idx == 0 else "",
                "Ед.изм.": ws_item.service.unit if ws_item else "",
                "Кол-во работы": float(ws_item.quantity) if ws_item else 0,
                "Цена работы": float(ws_item.unit_price) if ws_item else 0,
                "Сумма работы": float(ws_item.total_price) if ws_item else 0,
                "Сумма материалов": float(work.materials_total_price) if svc_idx == 0 and work.materials_total_price else 0,
                "ИТОГО": float(work.total_price) if svc_idx == 0 and work.total_price else 0,
            }

            if svc_idx == 0:
                for i, wm in enumerate(work.work_materials or []):
                    row[f"Материал {i+1}"] = wm.material.name
                    row[f"Кол-во {i+1}"] = float(wm.quantity)
                    row[f"Цена {i+1}"] = float(wm.unit_price)
                    row[f"Сумма {i+1}"] = float(wm.total_price)

                work_photos = work.photos or []
                for i, photo in enumerate(work_photos):
                    row[f"Фото {i+1}"] = photo.filename
                photo_map.append(work_photos)
            else:
                photo_map.append([])

            data.append(row)

    df = pd.DataFrame(data, columns=all_columns)

    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Работы')
        ws = writer.sheets['Работы']

        # Настройка колонок фото
        for i in range(max_photos):
            col_idx = len(base_columns) + len(material_columns) + i + 1  # 1-based
            col_letter = get_column_letter(col_idx)
            ws.column_dimensions[col_letter].width = 22

        # Вставка изображений
        for row_idx, photos in enumerate(photo_map, start=2):  # start=2 т.к. первая строка — заголовок
            if not photos:
                continue
            for photo_idx, photo in enumerate(photos):
                col_idx = len(base_columns) + len(material_columns) + photo_idx + 1
                col_letter = get_column_letter(col_idx)
                img_path = photo.file_path
                if img_path and os.path.exists(img_path):
                    try:
                        img = XLImage(img_path)
                        img.width = 160
                        img.height = 120
                        cell = f"{col_letter}{row_idx}"
                        ws.add_image(img, cell)
                        ws.row_dimensions[row_idx].height = 100
                    except Exception:
                        pass
    
    output.seek(0)
    
    filename = f"works_report_{date_from or 'all'}_{date_to or 'all'}.xlsx"
    
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


@router.get("/summary/export")
def export_summary(
    date_from: str = None,
    date_to: str = None,
    group_by: str = "building",
    db: Session = Depends(get_db),
    current_user = Depends(require_director)
):
    report = get_summary_report(date_from, date_to, group_by, None, db, current_user)
    
    df = pd.DataFrame(report["items"])
    
    # Добавление итогов
    totals = report["totals"]
    df.loc[len(df)] = {
        "group_name": "ИТОГО",
        "works_count": totals["works_count"],
        "service_total": totals["service_total"],
        "materials_total": totals["materials_total"],
        "total": totals["total"],
    }
    
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Сводный отчёт')
    
    output.seek(0)
    
    filename = f"works_summary_{group_by}_{date_from or 'all'}_{date_to or 'all'}.xlsx"
    
    from fastapi.responses import StreamingResponse
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


# === Акт сдачи-приемки в Word ===

def _plural_form(n: int, forms: List[str]) -> str:
    """Возвращает правильную форму слова для числа (forms: [0, 1, 2, 5])."""
    if n % 10 == 1 and n % 100 != 11:
        return forms[1]
    if 2 <= n % 10 <= 4 and (n % 100 < 10 or n % 100 >= 20):
        return forms[2]
    return forms[3]


def _convert_less_than_thousand(num: int, feminine: bool = False) -> str:
    """Преобразует число от 0 до 999 в текст."""
    if num == 0:
        return ""
    units_m = ["", "один", "два", "три", "четыре", "пять", "шесть", "семь", "восемь", "девять"]
    units_f = ["", "одна", "две", "три", "четыре", "пять", "шесть", "семь", "восемь", "девять"]
    teens = ["десять", "одиннадцать", "двенадцать", "тринадцать", "четырнадцать",
             "пятнадцать", "шестнадцать", "семнадцать", "восемнадцать", "девятнадцать"]
    tens = ["", "", "двадцать", "тридцать", "сорок", "пятьдесят", "шестьдесят",
            "семьдесят", "восемьдесят", "девяносто"]
    hundreds = ["", "сто", "двести", "триста", "четыреста", "пятьсот",
                "шестьсот", "семьсот", "восемьсот", "девятьсот"]
    units = units_f if feminine else units_m
    parts = []
    h = num // 100
    if h:
        parts.append(hundreds[h])
    t = (num % 100) // 10
    u = num % 10
    if t == 1:
        parts.append(teens[u])
    else:
        if t:
            parts.append(tens[t])
        if u:
            parts.append(units[u])
    return " ".join(parts)


def number_to_words_ru(n: int) -> str:
    """Преобразует целое положительное число в сумму прописью (рубли)."""
    if n < 0:
        return "минус " + number_to_words_ru(-n)
    if n == 0:
        return "ноль рублей"

    ruble_forms = ["", "рубль", "рубля", "рублей"]
    thousand_forms = ["", "тысяча", "тысячи", "тысяч"]
    million_forms = ["", "миллион", "миллиона", "миллионов"]
    billion_forms = ["", "миллиард", "миллиарда", "миллиардов"]

    parts = []
    remainder = n

    if remainder >= 1_000_000_000:
        b = remainder // 1_000_000_000
        parts.append(_convert_less_than_thousand(b) + " " + _plural_form(b, billion_forms))
        remainder %= 1_000_000_000

    if remainder >= 1_000_000:
        m = remainder // 1_000_000
        parts.append(_convert_less_than_thousand(m) + " " + _plural_form(m, million_forms))
        remainder %= 1_000_000

    if remainder >= 1_000:
        t = remainder // 1_000
        parts.append(_convert_less_than_thousand(t, feminine=True) + " " + _plural_form(t, thousand_forms))
        remainder %= 1_000

    if remainder > 0:
        parts.append(_convert_less_than_thousand(remainder) + " " + _plural_form(remainder, ruble_forms))
    elif parts:
        # если remainder == 0, но были тысячи/миллионы, добавляем "рублей"
        parts.append("рублей")

    return " ".join(parts).strip()


def _format_period_date(d: date) -> str:
    """Форматирует дату как '«15» июня 2026 г.'"""
    months = [
        "января", "февраля", "марта", "апреля", "мая", "июня",
        "июля", "августа", "сентября", "октября", "ноября", "декабря"
    ]
    return f"«{d.day:02d}» {months[d.month - 1]} {d.year} г."


def _format_request_label(work: Work) -> str:
    """Формирует подпись заявки вида '15 от 21.07.26'."""
    if not work.request_id:
        return ""
    request_date = work.request.created_at.strftime("%d.%m.%y") if work.request and work.request.created_at else ""
    if request_date:
        return f"{work.request_id} от {request_date}"
    return str(work.request_id)


def generate_act_docx(works: List[Work], date_from: str = None, date_to: str = None, contractor_name: str = "") -> BytesIO:
    """Генерирует Word-документ с актом сдачи-приемки по списку работ."""
    doc = Document()

    # Настройки страницы
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # Заголовок
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Акт сдачи-приемки оказанных услуг")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    # Город и дата
    today = datetime.now()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"г. Пятигорск\t{_format_period_date(today.date())}")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # Определяем период
    work_dates = [w.work_date for w in works if w.work_date]
    if work_dates:
        period_start = min(work_dates)
        period_end = max(work_dates)
    else:
        period_start = date.today()
        period_end = date.today()

    period_from_str = _format_period_date(period_start) if period_start else "«____» __________ 20__ г."
    period_to_str = _format_period_date(period_end) if period_end else "«____» __________ 20__ г."

    # Исполнитель
    contractor = contractor_name.strip() if contractor_name else "_________________________________________________________________________"

    # Основной текст
    def add_text(text: str, bold: bool = False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        if bold:
            run.bold = True
        p.paragraph_format.line_spacing = 1.15
        return p

    add_text(
        f"ФГАОУ ВО «Северо-Кавказский федеральный университет», именуемый в дальнейшем «Заказчик», "
        f"в лице ____________________________________________ от ___.___.20___г., и {contractor}, "
        f"именуемый в дальнейшем «Исполнитель», совместно именуемые «Стороны», составили настоящий акт о нижеследующем:"
    )

    add_text("1. Исполнитель оказал Заказчику услуги по договору возмездного оказания услуг № ___________ от «____» _____________ 20___г.")
    add_text(f"2. Исполнитель оказал услуги в срок: с {period_from_str} по {period_to_str}.")
    add_text("3. Исполнитель оказал услуги______________________________________________.")
    p = add_text("(своевременно / не своевременно)")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_text("4. Заказчик по объему и качеству оказанных Исполнителем услуг (претензии не имеет/претензии имеет)__________________________________________________________________________________________________________________________________________________________________________.")
    add_text("5. Подписанием Акта Стороны подтверждают выполнение условий договора возмездного оказания услуг ____________________________________________________________________________________.")
    p = add_text("(надлежащим образом/ с замечанием)")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Комендант(ы) из заявок
    comendants = set()
    for work in works:
        if work.request and work.request.creator and work.request.creator.role == "comendant":
            comendants.add(work.request.creator.full_name or work.request.creator.username)
    comendant_name = ", ".join(sorted(comendants)) if comendants else "_________________________________________________________________________"
    add_text(f"Комендант: {comendant_name}")

    # Таблица материалов
    add_text("6. Используемые материалы для оказания услуг:", bold=True)
    table = doc.add_table(rows=1, cols=9)
    table.style = "Table Grid"
    table.autofit = False
    table.allow_autofit = False

    hdr = table.rows[0].cells
    headers = ["№", "№ заявки", "Дата работ", "Адрес объекта", "Наименование материалов", "Кол-во", "Ед. измерения", "Цена за ед. руб.", "Стоимость материалов руб."]
    widths = [Cm(1.0), Cm(1.5), Cm(2.2), Cm(2.8), Cm(3.8), Cm(1.5), Cm(2.0), Cm(2.5), Cm(3.0)]
    for i, (h, w) in enumerate(zip(headers, widths)):
        hdr[i].text = h
        hdr[i].width = w
        for paragraph in hdr[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)

    materials_total = Decimal("0")
    row_num = 1
    for work in works:
        address = f"{work.building.number} {work.building.name or ''}".strip()
        work_date_str = work.work_date.strftime("%d.%m.%Y") if work.work_date else ""
        request_str = _format_request_label(work)
        for wm in work.work_materials or []:
            row = table.add_row().cells
            values = [
                str(row_num),
                request_str,
                work_date_str,
                address,
                wm.material.name,
                str(wm.quantity),
                wm.material.unit or "",
                f"{float(wm.unit_price):.2f}",
                f"{float(wm.total_price):.2f}",
            ]
            for i, val in enumerate(values):
                row[i].text = val
                row[i].width = widths[i]
            materials_total += wm.total_price or Decimal("0")
            row_num += 1

    # Добавляем пустые строки для визуального запаса
    while row_num <= 2:
        row = table.add_row().cells
        row[0].text = str(row_num)
        row_num += 1

    # Итого по материалам
    total_row = table.add_row().cells
    total_row[0].merge(total_row[7])
    total_row[0].text = "Итого:"
    total_row[8].text = f"{float(materials_total):.2f}"
    for paragraph in total_row[0].paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
    for paragraph in total_row[8].paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)

    # Таблица работ
    add_text("7. Объем оказанных услуг:", bold=True)
    table2 = doc.add_table(rows=1, cols=9)
    table2.style = "Table Grid"
    table2.autofit = False
    table2.allow_autofit = False

    hdr2 = table2.rows[0].cells
    headers2 = ["№", "№ заявки", "Дата работ", "Адрес объекта", "Наименование работ", "Кол-во", "Ед. измерения", "Цена за ед. руб.", "Стоимость работ (услуг) руб."]
    widths2 = [Cm(1.0), Cm(1.5), Cm(2.2), Cm(2.8), Cm(3.8), Cm(1.5), Cm(2.0), Cm(2.5), Cm(3.0)]
    for i, (h, w) in enumerate(zip(headers2, widths2)):
        hdr2[i].text = h
        hdr2[i].width = w
        for paragraph in hdr2[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)

    service_total = Decimal("0")
    row_num = 1
    for work in works:
        address = f"{work.building.number} {work.building.name or ''}".strip()
        work_date_str = work.work_date.strftime("%d.%m.%Y") if work.work_date else ""
        request_str = _format_request_label(work)
        for ws in work.work_services:
            row = table2.add_row().cells
            values = [
                str(row_num),
                request_str,
                work_date_str,
                address,
                ws.service.name,
                str(ws.quantity),
                ws.service.unit or "",
                f"{float(ws.unit_price or 0):.2f}",
                f"{float(ws.total_price or 0):.2f}",
            ]
            for i, val in enumerate(values):
                row[i].text = val
                row[i].width = widths2[i]
            service_total += ws.total_price or Decimal("0")
            row_num += 1

    for i in range(2 - (row_num - 1)):
        row = table2.add_row().cells
        row[0].text = str(row_num + i)

    total_row2 = table2.add_row().cells
    total_row2[0].merge(total_row2[7])
    total_row2[0].text = "Итого:"
    total_row2[8].text = f"{float(service_total):.2f}"
    for paragraph in total_row2[0].paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
    for paragraph in total_row2[8].paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)

    # Итого по акту
    total = sum((w.total_price or Decimal("0") for w in works), Decimal("0"))
    total_int = int(total)
    total_words = number_to_words_ru(total_int)
    add_text(f"8. Итого по акту оказания услуг и использованных материалов составляет: {float(total):.2f} ({total_words}) рублей, в том числе НДС, без НДС.")
    add_text("9. Акт составлен в двух экземплярах, имеющих равную юридическую силу, по одному для каждой из Сторон.")

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


@router.get("/act")
def export_act(
    date_from: str = None,
    date_to: str = None,
    building_id: str = None,
    service_id: str = None,
    user_id: str = None,
    work_id: str = None,
    db: Session = Depends(get_db),
    current_user = Depends(require_director)
):
    """Формирует Word-документ (акт сдачи-приемки) по отфильтрованным работам."""
    query = db.query(Work).options(
        joinedload(Work.request).joinedload(Request.creator),
    )

    if work_id:
        query = query.filter(Work.id == int(work_id))
    if date_from:
        query = query.filter(Work.work_date >= date_from)
    if date_to:
        query = query.filter(Work.work_date <= date_to)
    if building_id:
        query = query.filter(Work.building_id == int(building_id))
    if service_id:
        query = query.join(WorkService).filter(WorkService.service_id == int(service_id))
    if user_id:
        query = query.filter(Work.user_id == int(user_id))

    works = query.order_by(Work.work_date.desc()).all()

    if not works:
        raise HTTPException(status_code=404, detail="Нет данных для формирования акта")

    # Определяем исполнителя
    contractor_name = ""
    if user_id:
        user = db.query(User).filter(User.id == int(user_id)).first()
        contractor_name = user.full_name or user.username if user else ""
    else:
        users = set((w.user.full_name or w.user.username) for w in works)
        if len(users) == 1:
            contractor_name = list(users)[0]

    output = generate_act_docx(works, date_from, date_to, contractor_name)

    filename = f"act_{date_from or 'all'}_{date_to or 'all'}.docx"

    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )
